"""
PDF Content Extractor for Learning Management System
Extracts text from PDF/DOCX and uses AI to structure it into path, modules, and quizzes
"""
import os
import json
import re
from typing import Dict, List, Any, Optional
import google.generativeai as genai
from django.conf import settings

# Try to import PDF libraries
try:
    import PyPDF2
    PDF_LIBRARY = 'pypdf2'
except ImportError:
    try:
        import pdfplumber
        PDF_LIBRARY = 'pdfplumber'
    except ImportError:
        PDF_LIBRARY = None

# Try to import DOCX library
try:
    import docx
    DOCX_LIBRARY = True
except ImportError:
    DOCX_LIBRARY = False


def extract_text_from_pdf(pdf_file) -> str:
    """Extract text content from a PDF file"""
    if PDF_LIBRARY is None:
        raise ImportError("No PDF library available. Install PyPDF2 or pdfplumber.")
    
    text_content = []
    
    if PDF_LIBRARY == 'pypdf2':
        reader = PyPDF2.PdfReader(pdf_file)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_content.append(text)
    elif PDF_LIBRARY == 'pdfplumber':
        import pdfplumber
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
    
    return '\n\n'.join(text_content)


def extract_text_from_docx(docx_file) -> str:
    """Extract text content from a DOCX file"""
    if not DOCX_LIBRARY:
        raise ImportError("python-docx library not installed. Run: pip install python-docx")
    
    text_content = []
    document = docx.Document(docx_file)
    
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_content.append(paragraph.text)
    
    # Also extract text from tables
    for table in document.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_content.append(' | '.join(row_text))
    
    return '\n\n'.join(text_content)


def extract_text_from_file(file) -> str:
    """Extract text from PDF or DOCX file based on filename"""
    filename = file.name.lower() if hasattr(file, 'name') else ''
    
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(file)
    elif filename.endswith('.docx') or filename.endswith('.doc'):
        return extract_text_from_docx(file)
    else:
        # Try PDF first, then DOCX
        try:
            return extract_text_from_pdf(file)
        except:
            try:
                file.seek(0)  # Reset file pointer
                return extract_text_from_docx(file)
            except:
                raise ValueError("Could not extract text from file. Please upload a PDF or DOCX file.")


def parse_content_with_ai(text: str, content_type: str = 'full', model_type: str = None) -> Dict[str, Any]:
    """
    Use AI to parse extracted text into structured learning content.
    Now uses AIServiceFactory to support multiple AI models.
    
    Args:
        text: Extracted PDF text
        content_type: 'full' (path + modules + quiz), 'modules_only', 'quiz_only'
        model_type: AI model to use (google_gemini, openai_gpt4, anthropic_claude, local)
    
    Returns:
        Structured content dictionary
    """
    # Import AIServiceFactory from ai_mentor
    try:
        from apps.ai_mentor.services.ai_service import AIServiceFactory
    except ImportError:
        # Fallback to direct Gemini if ai_mentor not available
        return parse_content_with_gemini_direct(text, content_type)
    
    # Truncate text if too long (AI models have token limits)
    max_chars = 50000
    if len(text) > max_chars:
        text = text[:max_chars] + "\n\n[Content truncated...]"
    
    if content_type == 'full':
        prompt = create_full_extraction_prompt(text)
    elif content_type == 'modules_only':
        prompt = create_modules_extraction_prompt(text)
    elif content_type == 'quiz_only':
        prompt = create_quiz_extraction_prompt(text)
    else:
        prompt = create_full_extraction_prompt(text)
    
    try:
        # Get AI service based on model_type
        import logging
        logger = logging.getLogger(__name__)
        logger.info(f"parse_content_with_ai called with model_type: {model_type}")
        
        service = AIServiceFactory.get_service(model_type)
        logger.info(f"Got service: {service.__class__.__name__}")
        
        response_text = None
        try:
            response_text = service.generate_response(prompt)
        except Exception as primary_error:
            logger.warning(f"Primary model {model_type} failed: {primary_error}")
            
            # Try OpenRouter as fallback
            logger.info("Trying OpenRouter as fallback...")
            try:
                from apps.ai_mentor.services.ai_service import OpenRouterService
                fallback_service = OpenRouterService(model='google/gemini-2.0-flash-exp:free')
                logger.info(f"Using fallback: OpenRouterService with Gemini 2.0 Flash")
                response_text = fallback_service.generate_response(prompt)
            except Exception as fallback_error:
                logger.error(f"Fallback also failed: {fallback_error}")
                raise ValueError(f"AI processing failed: {primary_error}. Fallback also failed: {fallback_error}")
        
        if not response_text:
            raise ValueError("No response from AI service")
        
        # Log response for debugging
        logger.info(f"AI Response length: {len(response_text)} chars")
        logger.info(f"AI Response preview: {response_text[:300]}...")
        
        # Try multiple strategies to extract JSON
        parsed_json = None
        
        # Strategy 1: Look for JSON in markdown code blocks
        code_block_match = re.search(r'```(?:json)?\s*(\{[\s\S]*?\})\s*```', response_text)
        if code_block_match:
            try:
                parsed_json = json.loads(code_block_match.group(1))
                logger.info("Extracted JSON from markdown code block")
            except json.JSONDecodeError:
                pass
        
        # Strategy 2: Find outermost JSON object
        if not parsed_json:
            # Find the first { and last } to get the full JSON object
            first_brace = response_text.find('{')
            last_brace = response_text.rfind('}')
            if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
                json_str = response_text[first_brace:last_brace + 1]
                try:
                    parsed_json = json.loads(json_str)
                    logger.info("Extracted JSON using brace matching")
                except json.JSONDecodeError:
                    pass
        
        # Strategy 3: Use regex to find JSON-like structure
        if not parsed_json:
            json_match = re.search(r'\{[\s\S]*\}', response_text)
            if json_match:
                try:
                    parsed_json = json.loads(json_match.group())
                    logger.info("Extracted JSON using regex")
                except json.JSONDecodeError:
                    pass
        
        if parsed_json:
            return parsed_json
        else:
            # Log the response for debugging
            logger.error(f"Could not extract JSON from response: {response_text[:1000]}")
            raise ValueError(f"Could not extract JSON from AI response. Response preview: {response_text[:200]}...")
            
    except json.JSONDecodeError as e:
        raise ValueError(f"Failed to parse AI response as JSON: {e}")
    except ValueError:
        # Re-raise ValueErrors as-is
        raise
    except Exception as e:
        error_str = str(e)
        # Handle rate limit errors with helpful message
        if "429" in error_str or "quota" in error_str.lower() or "rate limit" in error_str.lower():
            raise ValueError(f"AI rate limit exceeded. Please wait a moment and try again, or switch to a different AI model in settings.")
        raise ValueError(f"AI processing failed: {e}")


def parse_content_with_gemini_direct(text: str, content_type: str = 'full') -> Dict[str, Any]:
    """
    Direct Gemini implementation as fallback.
    Used when AIServiceFactory is not available.
    """
    # Configure Gemini - check multiple possible env var names
    api_key = (
        getattr(settings, 'GEMINI_API_KEY', None) or 
        getattr(settings, 'GOOGLE_GEMINI_API_KEY', None) or
        os.environ.get('GEMINI_API_KEY') or
        os.environ.get('GOOGLE_GEMINI_API_KEY')
    )
    if not api_key:
        raise ValueError("GEMINI_API_KEY or GOOGLE_GEMINI_API_KEY not configured in settings or environment")
    
    genai.configure(api_key=api_key)
    
    # Model names to try
    model_names = [
        'models/gemini-2.5-flash',
        'models/gemini-flash-latest',
        'models/gemini-2.0-flash',
        'models/gemini-pro-latest',
    ]
    model = None
    last_error = None
    
    for model_name in model_names:
        try:
            model = genai.GenerativeModel(model_name)
            break
        except Exception as e:
            last_error = e
            continue
    
    if model is None:
        raise ValueError(f"No available Gemini model found. Tried: {model_names}. Last error: {last_error}")
    
    # Truncate text if too long
    max_chars = 50000
    if len(text) > max_chars:
        text = text[:max_chars] + "\n\n[Content truncated...]"
    
    if content_type == 'full':
        prompt = create_full_extraction_prompt(text)
    elif content_type == 'modules_only':
        prompt = create_modules_extraction_prompt(text)
    elif content_type == 'quiz_only':
        prompt = create_quiz_extraction_prompt(text)
    else:
        prompt = create_full_extraction_prompt(text)
    
    try:
        response = model.generate_content(prompt)
        result_text = response.text
        
        # Extract JSON from response
        json_match = re.search(r'\{[\s\S]*\}', result_text)
        if json_match:
            return json.loads(json_match.group())
        else:
            raise ValueError("Could not extract JSON from AI response")
            
    except json.JSONDecodeError as e:
        raise ValueError(f"Failed to parse AI response as JSON: {e}")
    except Exception as e:
        raise ValueError(f"AI processing failed: {e}")


def create_full_extraction_prompt(text: str) -> str:
    """Create prompt for full content extraction (path + modules + quizzes)"""
    return f"""Analyze this educational content and extract structured learning materials.

CONTENT:
{text}

Extract and return a JSON object with the following structure:
{{
    "path": {{
        "name": "Course/Path title (extract from content or generate appropriate title)",
        "description": "2-3 sentence description of the course",
        "program_type": "bsit" or "bscs" or "bsis" (infer from content, default to "bsit"),
        "difficulty_level": "beginner" or "intermediate" or "advanced",
        "estimated_duration": number (weeks, estimate based on content volume),
        "required_skills": ["skill1", "skill2"] (list of prerequisites)
    }},
    "modules": [
        {{
            "title": "Module title",
            "description": "Brief module description",
            "content": "Full module content in HTML format with proper formatting",
            "module_type": "text" or "video" or "interactive",
            "difficulty_level": "beginner" or "intermediate" or "advanced",
            "duration_minutes": number (estimate reading/learning time),
            "order": 1
        }}
    ],
    "quizzes": [
        {{
            "module_index": 0,
            "title": "Quiz title",
            "description": "Quiz description",
            "questions": [
                {{
                    "question_text": "The question",
                    "question_type": "multiple_choice" or "true_false",
                    "choices": ["Actual answer option from content", "Another real answer option", "A third real option", "A fourth real option"],
                    "correct_answer": "The correct option text or index",
                    "explanation": "Why this is correct",
                    "points": 1
                }}
            ]
        }}
    ]
}}

IMPORTANT RULES:
1. Split content into logical modules (chapters, sections, topics)
2. Each module should be self-contained learning unit
3. Generate 3-5 quiz questions per module based on key concepts
4. Use HTML formatting for module content (<h2>, <p>, <ul>, <li>, <strong>, <code>)
5. Ensure questions test understanding, not just memorization
6. If content doesn't have clear sections, create logical divisions
7. Return ONLY valid JSON, no other text
8. CRITICAL: Quiz choices MUST contain ACTUAL answer text from the document. NEVER use generic placeholders like "Option A", "Option B", etc.

Return the JSON:"""


def create_modules_extraction_prompt(text: str) -> str:
    """Create prompt for modules-only extraction"""
    return f"""Analyze this educational content and extract learning modules.

CONTENT:
{text}

Extract and return a JSON object with modules array:
{{
    "modules": [
        {{
            "title": "Module title",
            "description": "Brief module description",
            "content": "Full module content in HTML format",
            "module_type": "text",
            "difficulty_level": "beginner" or "intermediate" or "advanced",
            "duration_minutes": number,
            "order": 1
        }}
    ]
}}

Split content into logical learning modules. Use HTML formatting.
Return ONLY valid JSON:"""


def create_quiz_extraction_prompt(text: str) -> str:
    """Create prompt for quiz-only extraction with frontend-compatible format"""
    return f"""Analyze this educational content and generate quiz questions.

CONTENT:
{text}

Generate quiz questions and return a JSON object with this EXACT structure:
{{
    "questions": [
        {{
            "id": "unique_string_id",
            "title": "Short question title",
            "content": "Full question text with any context or code examples in HTML format",
            "type": "multiple_choice",
            "choices": [
                {{"id": "1", "text": "The actual first answer option extracted from content", "isCorrect": false}},
                {{"id": "2", "text": "The actual second answer option extracted from content", "isCorrect": true}},
                {{"id": "3", "text": "The actual third answer option extracted from content", "isCorrect": false}},
                {{"id": "4", "text": "The actual fourth answer option extracted from content", "isCorrect": false}}
            ],
            "correctAnswer": null,
            "points": 1
        }},
        {{
            "id": "unique_string_id",
            "title": "True/False question title",
            "content": "Question content in HTML",
            "type": "true_false",
            "choices": null,
            "correctAnswer": "true",
            "points": 1
        }},
        {{
            "id": "unique_string_id",
            "title": "Enumeration question title",
            "content": "List the items...",
            "type": "enumeration",
            "choices": null,
            "correctAnswer": "item1, item2, item3",
            "points": 2
        }}
    ]
}}

QUESTION TYPES TO USE:
- "multiple_choice": Has choices array with isCorrect flag (ONLY ONE should be true)
- "true_false": correctAnswer is either "true" or "false"
- "enumeration": correctAnswer is comma-separated list of items
- "short_answer": correctAnswer is the expected text
- "essay": No correctAnswer needed (manual grading)

RULES:
1. Generate 8-15 questions based on key concepts
2. Use HTML in content field for formatting (<p>, <code>, <strong>, <ul>, <li>)
3. For multiple_choice, ensure exactly ONE choice has isCorrect: true
4. Mix question types appropriately based on content
5. Assign points: 1 for simple, 2 for medium, 3 for complex questions
6. Each id must be unique (use numbers as strings: "1", "2", etc.)
7. Return ONLY valid JSON, no other text
8. CRITICAL: Each choice "text" MUST contain the ACTUAL answer text from the document content. NEVER use generic placeholders like "Option A", "Option B", "Option C", "Option D". Every choice must be a real, meaningful answer option."""


def process_pdf_for_learning(pdf_file, extraction_type: str = 'full', model_type: str = None) -> Dict[str, Any]:
    """
    Main function to process PDF/DOCX and extract learning content
    
    Args:
        pdf_file: Uploaded PDF or DOCX file
        extraction_type: 'full', 'modules_only', 'quiz_only'
        model_type: AI model to use (google_gemini, openai_gpt4, anthropic_claude, local)
    
    Returns:
        Structured learning content
    """
    # Step 1: Extract text from PDF or DOCX
    text = extract_text_from_file(pdf_file)
    
    if not text or len(text.strip()) < 100:
        raise ValueError("File appears to be empty or contains too little text")
    
    # Step 2: Parse with AI (using specified or default model)
    structured_content = parse_content_with_ai(text, extraction_type, model_type)
    
    # Step 3: Validate and clean the structure
    validated_content = validate_extracted_content(structured_content, extraction_type)
    
    return validated_content


def validate_extracted_content(content: Dict[str, Any], extraction_type: str) -> Dict[str, Any]:
    """Validate and clean extracted content"""
    
    if extraction_type == 'full':
        # Validate path
        if 'path' not in content:
            content['path'] = {
                'name': 'Untitled Course',
                'description': 'Course extracted from PDF',
                'program_type': 'bsit',
                'difficulty_level': 'beginner',
                'estimated_duration': 4,
                'required_skills': []
            }
        
        # Validate modules
        if 'modules' not in content or not content['modules']:
            content['modules'] = [{
                'title': 'Module 1',
                'description': 'Content from PDF',
                'content': '<p>Content will be added</p>',
                'module_type': 'text',
                'difficulty_level': 'beginner',
                'duration_minutes': 30,
                'order': 1
            }]
        
        # Ensure order is set
        for i, module in enumerate(content['modules']):
            module['order'] = i + 1
            if 'duration_minutes' not in module:
                module['duration_minutes'] = 30
        
        # Validate quizzes
        if 'quizzes' not in content:
            content['quizzes'] = []
    
    elif extraction_type == 'modules_only':
        if 'modules' not in content:
            content['modules'] = []
    
    elif extraction_type == 'quiz_only':
        if 'questions' not in content:
            content['questions'] = []
    
    return content
